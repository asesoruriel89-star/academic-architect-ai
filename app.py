import os, time, re, io
from flask import Flask, render_template, request, send_file, jsonify
from groq import Groq
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from dotenv import load_dotenv

# Carga las llaves desde el archivo .env
load_dotenv()

app = Flask(__name__)

# Configuración de Multi-Keys desde entorno
keys_raw = os.getenv("GROQ_API_KEYS")
LISTA_API_KEYS = keys_raw.split(",") if keys_raw else []

MODELO = "openai/gpt-oss-120b"
INDICE_KEY_ACTUAL = 0

def obtener_cliente():
    global INDICE_KEY_ACTUAL
    return Groq(api_key=LISTA_API_KEYS[INDICE_KEY_ACTUAL])

def rotar_key():
    global INDICE_KEY_ACTUAL
    INDICE_KEY_ACTUAL = (INDICE_KEY_ACTUAL + 1) % len(LISTA_API_KEYS)

def limpiar_texto_markdown(texto):
    texto = re.sub(r'#+\s*.*?\n', '', texto) 
    texto = re.sub(r'\*\*(.*?)\*\*', r'\1', texto)
    texto = re.sub(r'\*(.*?)\*', r'\1', texto)
    texto = re.sub(r'^\s*[-•]\s+', '', texto, flags=re.MULTILINE)
    return texto.strip()

def llamar_ai(prompt, sys_msg):
    global INDICE_KEY_ACTUAL
    try:
        client = obtener_cliente()
        resp = client.chat.completions.create(
            model=MODELO,
            messages=[{"role": "system", "content": sys_msg}, {"role": "user", "content": prompt}],
            temperature=0.5
        )
        return resp.choices[0].message.content
    except Exception:
        rotar_key()
        time.sleep(2)
        return llamar_ai(prompt, sys_msg)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/get_structure', methods=['POST'])
def get_structure():
    data = request.json
    topic = data.get('topic')
    doc_type = data.get('type')
    count = 35 if doc_type == "Thesis" else 15 if doc_type == "Book Chapter" else 6
    sys_msg = "You are an Academic Director. Design a professional research structure."
    prompt = f"Create a list of exactly {count} sub-topic titles for a professional {doc_type} about '{topic}'. Return ONLY titles, one per line."
    raw_plan = llamar_ai(prompt, sys_msg)
    plan = [line.strip() for line in raw_plan.split('\n') if len(line.strip()) > 5]
    return jsonify({"plan": plan[:count]})

@app.route('/generate_content', methods=['POST'])
def generate_content():
    data = request.json
    topic = data.get('topic')
    section = data.get('section')
    sys_msg = "PhD Researcher. Write 800 words, no markdown, APA 7 citations."
    prompt = f"Write the academic development for: '{section}' about '{topic}'."
    content = llamar_ai(prompt, sys_msg)
    return jsonify({"content": limpiar_texto_markdown(content)})

@app.route('/export_docx', methods=['POST'])
def export_docx():
    data = request.json
    topic = data.get('topic')
    sections = data.get('sections')
    doc = Document()
    for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n\n\n\n{topic.upper()}\n\nACADEMIC RESEARCH\n")
    run.font.name = 'Times New Roman'; run.font.size = Pt(20); run.font.bold = True
    doc.add_page_break()
    for item in sections:
        h = doc.add_heading(item['title'], level=1)
        for r in h.runs: r.font.color.rgb = RGBColor(0,0,0); r.font.name = 'Times New Roman'
        para = doc.add_paragraph(item['text'])
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for r in para.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        doc.add_page_break()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="Research.docx")

if __name__ == '__main__':
    app.run(debug=True)